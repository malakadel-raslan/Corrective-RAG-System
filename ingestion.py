"""
Ingestion pipeline for the Corrective RAG system.

Handles: upload -> extract -> clean -> split text
Supported formats: PDF, DOCX, CSV, TXT
"""
import os
import re
from dataclasses import dataclass
from typing import List

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter


@dataclass
class RawDocument:
    text: str
    source: str  # filename, used later for citation


# --------------------------------------------------------------------------
# Extraction
# --------------------------------------------------------------------------

def extract_pdf(path: str) -> str:
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    return "\n".join(pages)


def extract_docx(path: str) -> str:
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs]
    # also pull table cell text, docs often store specs/data in tables
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def extract_csv(path: str) -> str:
    df = pd.read_csv(path)
    # flatten each row into a readable "column: value" line so it chunks
    # into meaningful, self-contained retrieval units
    lines = []
    for _, row in df.iterrows():
        line = ", ".join(f"{col}: {row[col]}" for col in df.columns)
        lines.append(line)
    return "\n".join(lines)


def extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".csv": extract_csv,
    ".txt": extract_txt,
}


def extract_file(path: str) -> RawDocument:
    ext = os.path.splitext(path)[1].lower()
    if ext not in EXTRACTORS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {list(EXTRACTORS.keys())}"
        )
    text = EXTRACTORS[ext](path)
    return RawDocument(text=text, source=os.path.basename(path))


# --------------------------------------------------------------------------
# Cleaning
# --------------------------------------------------------------------------

def clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --------------------------------------------------------------------------
# Chunking
# --------------------------------------------------------------------------

def chunk_documents(
    raw_docs: List[RawDocument],
    chunk_size: int = 800,
    chunk_overlap: int = 120,
):
    """
    Returns a list of langchain Document objects, each carrying
    metadata={"source": filename, "chunk_id": int} for later citation.
    """
    from langchain_core.documents import Document as LCDocument

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    lc_documents = []
    for raw in raw_docs:
        cleaned = clean_text(raw.text)
        if not cleaned:
            continue
        chunks = splitter.split_text(cleaned)
        for i, chunk in enumerate(chunks):
            lc_documents.append(
                LCDocument(
                    page_content=chunk,
                    metadata={"source": raw.source, "chunk_id": i},
                )
            )
    return lc_documents


def process_uploaded_files(paths: List[str]):
    """End-to-end: paths -> extracted -> cleaned -> chunked LC documents."""
    raw_docs = [extract_file(p) for p in paths]
    return chunk_documents(raw_docs)
